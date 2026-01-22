# src/routes/notices.py
from flask import Blueprint, jsonify, request, send_file, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from src.models.user import User
from datetime import datetime
import io
import os
import re
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib import colors
from PyPDF2 import PdfReader, PdfWriter
from docx import Document

notices_bp = Blueprint('notices', __name__)


def generate_rough_sleeper_docx(notice_data):
    """
    Generate a Rough Sleeper Notice using the Word template.
    Replaces placeholder text with actual form data.
    Returns bytes of the .docx file.
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'rough_sleeper_template.docx')

    # Load the template
    doc = Document(template_path)

    # Get form data with defaults
    property_address = notice_data.get('property_address', '[PROPERTY ADDRESS]')
    landowner_name = notice_data.get('landowner_name', '[LANDOWNER NAME]')

    # Format dates to UK format
    vacate_date_raw = notice_data.get('vacate_date', '')
    if vacate_date_raw:
        try:
            dt = datetime.strptime(vacate_date_raw, '%Y-%m-%d')
            vacate_date = dt.strftime('%d/%m/%Y')
        except ValueError:
            vacate_date = vacate_date_raw
    else:
        vacate_date = '[DATE]'

    vacate_time = notice_data.get('vacate_time', '12:00')

    date_served_raw = notice_data.get('date_served', '')
    if date_served_raw:
        try:
            dt = datetime.strptime(date_served_raw, '%Y-%m-%d')
            date_served = dt.strftime('%d/%m/%Y')
        except ValueError:
            date_served = date_served_raw
    else:
        date_served = datetime.now().strftime('%d/%m/%Y')

    # Define replacement mappings - template text -> new value
    replacements = {
        # Property address (the example in template)
        '11 Berkshire Road Camberley Surrey GU15 4DG': property_address,
        # Landowner name (the example in template)
        'CIVIC LTD': landowner_name,
        # Vacate deadline (the example in template)
        '12:00-23/01/2026': f'{vacate_time}-{vacate_date}',
        # Date served (the example in template)
        '22/01/2026': date_served,
    }

    # Replace text in all paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                # Preserve formatting by replacing in runs
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                # Also check if text is split across runs
                if old_text in para.text:
                    # Full paragraph replacement as fallback
                    inline = para.runs
                    full_text = para.text
                    new_full_text = full_text.replace(old_text, new_text)
                    if full_text != new_full_text:
                        # Clear and rewrite
                        for run in inline:
                            run.text = ''
                        if inline:
                            inline[0].text = new_full_text

    # Replace text in tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

    # Save to buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer.getvalue()


def require_admin():
    """Helper to check if current user is admin"""
    user_id = get_jwt_identity()
    user = User.query.get(user_id)
    if not user or user.role != 'admin':
        return None
    return user

def generate_notice_pdf(notice_type, notice_data):
    """
    Generate a notice PDF using the same template system as V3 job reports.
    Returns bytes of the PDF file.
    """
    # Path to the headed template
    template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'headed_template.pdf')
    
    # First, generate the content PDF
    content_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        content_buffer,
        pagesize=A4,
        rightMargin=1.5*cm,
        leftMargin=1.5*cm,
        topMargin=2.4*cm,  # REDUCED - fit on one page
        bottomMargin=1.0*cm  # REDUCED - fit on one page
    )
    
    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='NoticeTitle',
        parent=styles['Heading1'],
        fontSize=16,  # REDUCED from 20
        spaceAfter=12,  # REDUCED from 20
        alignment=TA_CENTER,
        textColor=colors.HexColor('#ff0000'),  # Red for legal notices
        underline=1  # ADD UNDERLINE
    ))
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading2'],
        fontSize=11,  # REDUCED from 14
        spaceBefore=8,  # REDUCED from 15
        spaceAfter=6,  # REDUCED from 10
        textColor=colors.HexColor('#ff6b35')
    ))
    styles.add(ParagraphStyle(
        name='FieldLabel',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        fontName='Helvetica-Bold'
    ))
    styles.add(ParagraphStyle(
        name='FieldValue',
        parent=styles['Normal'],
        fontSize=9,  # REDUCED from 11
        textColor=colors.HexColor('#1a1a2e'),
        spaceAfter=4  # REDUCED from 8
    ))
    styles.add(ParagraphStyle(
        name='LegalText',
        parent=styles['Normal'],
        fontSize=8.5,  # REDUCED from 10
        textColor=colors.HexColor('#333333'),
        spaceBefore=2  # REDUCED from 5
    ))
    
    elements = []
    
    if notice_type == 'notice_to_vacate':
        elements = generate_notice_to_vacate_content(notice_data, styles)
    elif notice_type == 'abandoned_vehicle':
        elements = generate_abandoned_vehicle_content(notice_data, styles)
    elif notice_type == 'rough_sleeper':
        elements = generate_rough_sleeper_content(notice_data, styles)
    else:
        raise ValueError('Invalid notice type')
    
    # Build content PDF
    doc.build(elements)
    content_buffer.seek(0)
    
    # Now merge content with the headed template
    try:
        # Read the template
        template_reader = PdfReader(template_path)
        template_page = template_reader.pages[0]
        
        # Read the content we just generated
        content_reader = PdfReader(content_buffer)
        
        # Create output PDF
        output = PdfWriter()
        
        # For each page of content, merge with template
        for i, content_page in enumerate(content_reader.pages):
            # Create a copy of the template for each page
            from copy import copy
            new_page = copy(template_page)
            # Merge the content on top of the template
            new_page.merge_page(content_page)
            output.add_page(new_page)
        
        # Write to buffer
        output_buffer = io.BytesIO()
        output.write(output_buffer)
        output_buffer.seek(0)
        return output_buffer.getvalue()
    
    except Exception as merge_error:
        # If merging fails, fall back to content-only PDF
        current_app.logger.warning(f"Template merge failed, using content only: {str(merge_error)}")
        content_buffer.seek(0)
        return content_buffer.getvalue()

def generate_notice_to_vacate_content(data, styles):
    """Generate content for Notice to Vacate"""
    elements = []
    
    # Title - smaller font
    elements.append(Paragraph("LEGAL NOTICE TO VACATE PREMISES", styles['NoticeTitle']))
    elements.append(Spacer(1, 0.08*inch))  # REDUCED to fit on one page
    
    # Property Address
    elements.append(Paragraph("PROPERTY ADDRESS", styles['SectionTitle']))
    address_text = data.get('property_address', '[ADDRESS NOT PROVIDED]')
    elements.append(Paragraph(address_text, styles['FieldValue']))
    elements.append(Spacer(1, 0.05*inch))  # REDUCED to fit one page
    
    # Date
    notice_date = data.get('date', datetime.now().strftime('%d/%m/%Y'))  # Changed to DD/MM/YYYY
    date_data = [['Date:', notice_date]]
    date_table = Table(date_data, colWidths=[0.6*inch, 1.5*inch])  # Tightened - reduced gap
    date_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),  # REDUCED to fit one page
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),  # REDUCED to fit one page
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # Left align
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),  # Left align date value
    ]))
    elements.append(date_table)
    elements.append(Spacer(1, 0.05*inch))  # REDUCED to fit one page
    
    # To Occupiers
    elements.append(Paragraph("TO: THE OCCUPIERS", styles['SectionTitle']))
    client_name = data.get('client_name', 'the legal owner')
    intro_text = f"We are writing on behalf of {client_name} of the above property."
    elements.append(Paragraph(intro_text, styles['FieldValue']))
    elements.append(Spacer(1, 0.05*inch))  # REDUCED to fit one page
    
    # Take Notice
    elements.append(Paragraph("TAKE NOTICE:", styles['SectionTitle']))
    
    # Compact notice points - combine some to save space
    notices = [
        "1. You are currently occupying the above premises WITHOUT the permission or authority of the legal owner. Your occupation constitutes TRESPASS under English Law.",
        "",
        "2. The legal owner has NOT granted you any right, licence, or permission to occupy these premises.",
        "",
        "3. You are hereby required to VACATE THE PREMISES IMMEDIATELY and remove all of your belongings and any persons under your control.",
        "",
        "4. V3 Services Ltd has been instructed by the legal owner to take all lawful steps necessary to secure possession of this property.",
        "",
        "5. Failure to vacate will result in the legal owner pursuing formal possession proceedings through the County Court, which may result in a Court Order for possession and costs being awarded against you.",
        "",
        "6. Any damage to the property, theft, or interference with utilities will be reported to the police and may result in criminal prosecution.",
    ]
    
    for notice in notices:
        if notice:
            elements.append(Paragraph(notice, styles['LegalText']))
        else:
            elements.append(Spacer(1, 0.02*inch))  # REDUCED to fit one page
    
    elements.append(Spacer(1, 0.05*inch))  # REDUCED to fit one page
    
    # Legal Warning - more compact
    elements.append(Paragraph("LEGAL WARNING", styles['SectionTitle']))
    warning_text = """Under Section 144 of the Legal Aid, Sentencing and Punishment of Offenders Act 2012, 
    you are committing an act of trespass. The property owner is entitled to take lawful action to recover possession. 
    This notice serves as formal notification that you must leave immediately. The property owner reserves all legal rights."""
    elements.append(Paragraph(warning_text, styles['LegalText']))
    elements.append(Spacer(1, 0.08*inch))  # REDUCED to fit one page
    
    # Issued By - compact table (NO DIRECTOR LINE)
    issued_data = [
        ['Issued by:', 'V3 Services Ltd'],
        ['Contact:', data.get('contact_phone', '0203 576 1343')],
        ['Email:', data.get('contact_email', 'Info@V3-Services.com')],
    ]
    
    issued_table = Table(issued_data, colWidths=[1.5*inch, 5*inch])
    issued_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),  # REDUCED to fit one page
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),  # REDUCED to fit one page
    ]))
    elements.append(issued_table)
    
    return elements

def generate_abandoned_vehicle_content(data, styles):
    """Generate content for Abandoned Vehicle Report"""
    elements = []
    
    # Title
    elements.append(Paragraph("ABANDONED VEHICLE REPORT", styles['NoticeTitle']))
    elements.append(Spacer(1, 0.3*inch))
    
    # Report Details
    elements.append(Paragraph("REPORT DETAILS", styles['SectionTitle']))
    report_data = [
        ['Date Reported:', data.get('date', datetime.now().strftime('%d/%m/%Y'))],  # Changed to DD/MM/YYYY
        ['Location:', data.get('location', '[LOCATION NOT PROVIDED]')],
        ['Client:', data.get('client_name', '[CLIENT NAME NOT PROVIDED]')],
    ]
    
    report_table = Table(report_data, colWidths=[2*inch, 4.5*inch])
    report_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(report_table)
    elements.append(Spacer(1, 0.2*inch))
    
    # Vehicle Details
    elements.append(Paragraph("VEHICLE DETAILS", styles['SectionTitle']))
    vehicle_data = [
        ['Registration:', data.get('registration', '[REG NOT PROVIDED]')],
        ['Make/Model:', data.get('make_model', '[NOT PROVIDED]')],
        ['Colour:', data.get('colour', '[NOT PROVIDED]')],
        ['Condition:', data.get('condition', '[NOT PROVIDED]')],
    ]
    
    vehicle_table = Table(vehicle_data, colWidths=[2*inch, 4.5*inch])
    vehicle_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(vehicle_table)
    elements.append(Spacer(1, 0.2*inch))
    
    # Description
    elements.append(Paragraph("DESCRIPTION", styles['SectionTitle']))
    description = data.get('description', 'Vehicle appears to have been abandoned on the property.')
    elements.append(Paragraph(description, styles['FieldValue']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Action Taken
    elements.append(Paragraph("ACTION TAKEN", styles['SectionTitle']))
    actions = [
        "• Vehicle documented with photographs",
        "• DVLA lookup completed",
        "• Notice affixed to vehicle",
        f"• Property owner notified: {data.get('client_name', '[CLIENT]')}",
    ]
    for action in actions:
        elements.append(Paragraph(action, styles['LegalText']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Recommendations
    elements.append(Paragraph("RECOMMENDATIONS", styles['SectionTitle']))
    recommendations = [
        "• Property owner to contact local authority for removal",
        "• Continue monitoring for 7 days",
        "• If not removed, escalate to legal action",
    ]
    for rec in recommendations:
        elements.append(Paragraph(rec, styles['LegalText']))
    elements.append(Spacer(1, 0.3*inch))
    
    # Report By
    report_by_data = [
        ['Report compiled by:', 'V3 Services Ltd'],
        ['Agent:', data.get('agent_name', '[AGENT NAME]')],
        ['Contact:', '0203 576 1343'],
    ]
    
    report_by_table = Table(report_by_data, colWidths=[2*inch, 4.5*inch])
    report_by_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(report_by_table)

    return elements


def generate_rough_sleeper_content(data, styles):
    """Generate content for Rough Sleeper Notice - 2 pages"""
    elements = []

    # Custom styles for rough sleeper notice
    styles.add(ParagraphStyle(
        name='RSTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=16,
        spaceBefore=8,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#cc0000'),  # Red color
        fontName='Helvetica-Bold',
        underline=1  # Underlined
    ))
    styles.add(ParagraphStyle(
        name='RSSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=8,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#333333'),
        fontName='Helvetica-Bold'
    ))
    styles.add(ParagraphStyle(
        name='RSLegalText',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#333333'),
        spaceBefore=6,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        leading=13
    ))
    styles.add(ParagraphStyle(
        name='RSBoldText',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#1a1a2e'),
        fontName='Helvetica-Bold',
        spaceBefore=8,
        spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        name='RSDeadline',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#cc0000'),  # Red color
        fontName='Helvetica-Bold',
        alignment=TA_CENTER,
        spaceBefore=12,
        spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        name='RSCharityTitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=10,
        spaceAfter=8,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#ff6b35')
    ))
    styles.add(ParagraphStyle(
        name='RSCharityText',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#333333'),
        alignment=TA_CENTER,
        spaceAfter=4
    ))

    # ============ PAGE 1: LEGAL NOTICE ============

    # Add spacer to push content down from header
    elements.append(Spacer(1, 1.2*inch))

    # Header - "To Persons Unknown and Belongings" with property address
    property_address = data.get('property_address', '[PROPERTY ADDRESS]')
    landowner_name = data.get('landowner_name', '[LANDOWNER NAME]')

    header_text = f'To Persons Unknown and Belongings ("the Belongings") at<br/>{property_address}'
    elements.append(Paragraph(header_text, styles['RSSubtitle']))
    elements.append(Spacer(1, 0.25*inch))

    # Title - Red and underlined
    elements.append(Paragraph('<u>NOTICE TO VACATE PRIVATE PROPERTY</u>', styles['RSTitle']))
    elements.append(Spacer(1, 0.15*inch))

    # Main Notice Text - landowner intro
    intro_text = f"""<b>{landowner_name} ("the Landowner")</b> is the Landowner of the land on which you are
    trespassing. You are here without permission, license or consent, and have no right to
    remain. Therefore, we are giving you notice that the Landowner requires you to leave. You
    must vacate this land with all your belongings by:"""
    elements.append(Paragraph(intro_text, styles['RSLegalText']))

    # Vacate Deadline - Red, centered, bold - format to UK date
    vacate_date_raw = data.get('vacate_date', '')
    if vacate_date_raw:
        try:
            # Parse ISO format and convert to UK format
            dt = datetime.strptime(vacate_date_raw, '%Y-%m-%d')
            vacate_date = dt.strftime('%d/%m/%Y')
        except ValueError:
            vacate_date = vacate_date_raw
    else:
        vacate_date = '[DATE]'
    vacate_time = data.get('vacate_time', '12:00')
    deadline_text = f'<u>{vacate_time} – {vacate_date}</u>'
    elements.append(Paragraph(deadline_text, styles['RSDeadline']))
    elements.append(Spacer(1, 0.1*inch))

    # Self-help warning (in red)
    selfhelp_text = """If you fail to vacate site and remove the Belongings, the Landowner will exercise its right of
    self-help to remove you and the Belongings from the land, even in your absence, and have
    Enforcement Agents to do so on its behalf."""
    styles.add(ParagraphStyle(
        name='RSRedText',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#cc0000'),
        spaceBefore=6,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        leading=13
    ))
    elements.append(Paragraph(selfhelp_text, styles['RSRedText']))
    elements.append(Spacer(1, 0.1*inch))

    # Torts Act section (in red)
    torts_text = """Please note, under Section 12 of the Torts (Interference with Goods) Act 1977, the
    Landowner intends to dispose of and/or destroy any belongings left on the site unless they
    are removed by the time and date specified above."""
    elements.append(Paragraph(torts_text, styles['RSRedText']))
    elements.append(Spacer(1, 0.1*inch))

    # Abandoned belongings warning (bold black)
    abandoned_text = """<b>ANY BELONGINGS LEFT ON SITE AFTER THIS TIME WILL BE DEEMED ABANDONED.</b>"""
    elements.append(Paragraph(abandoned_text, styles['RSBoldText']))
    elements.append(Spacer(1, 0.1*inch))

    # Liability Disclaimer (bold black, all caps style)
    disclaimer_text = """<b>PLEASE NOTE THAT THE LANDOWNER WILL NOT BE LIABLE FOR ANY ACTUAL OR
    CONSEQUENTIAL LOSS SUFFERED BY YOU OR ANY THIRD PARTY AS A RESULT OF
    DISPOSING OR DESTROYING SAID BELONGINGS.</b>"""
    elements.append(Paragraph(disclaimer_text, styles['RSBoldText']))
    elements.append(Spacer(1, 0.25*inch))

    # Signature Section - matching competitor layout
    assets_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'static', 'notice-assets')
    signature_path = os.path.join(assets_dir, 'signature.png')

    # Signature with image
    if os.path.exists(signature_path):
        try:
            sig_img = Image(signature_path, width=1.5*inch, height=0.6*inch)
            sig_table = Table([['Signed:', sig_img]], colWidths=[0.8*inch, 2*inch])
            sig_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                ('FONTSIZE', (0, 0), (0, 0), 10),
            ]))
            elements.append(sig_table)
        except Exception:
            elements.append(Paragraph("Signed: <i>V3 Services</i>", styles['RSLegalText']))
    else:
        elements.append(Paragraph("Signed: <i>V3 Services</i>", styles['RSLegalText']))

    elements.append(Paragraph("<b>On behalf of the Landowner</b>", styles['RSBoldText']))

    # Date Served - format to UK date
    date_served_raw = data.get('date_served', '')
    if date_served_raw:
        try:
            # Parse ISO format and convert to UK format
            dt = datetime.strptime(date_served_raw, '%Y-%m-%d')
            date_served = dt.strftime('%d/%m/%Y')
        except ValueError:
            date_served = date_served_raw
    else:
        date_served = datetime.now().strftime('%d/%m/%Y')
    elements.append(Paragraph(f"Dated: {date_served}", styles['RSLegalText']))
    elements.append(Spacer(1, 0.15*inch))

    # V3 Stamp image
    stamp_path = os.path.join(assets_dir, 'v3-stamp.png')
    if os.path.exists(stamp_path):
        try:
            stamp_img = Image(stamp_path, width=1.0*inch, height=1.0*inch)
            elements.append(stamp_img)
        except Exception as e:
            current_app.logger.warning(f"Could not load stamp image: {e}")

    elements.append(Spacer(1, 0.2*inch))

    # Contact details section
    elements.append(Paragraph("<i>Any queries in relation to this notice should be sent in writing to:</i>", styles['RSLegalText']))
    elements.append(Spacer(1, 0.15*inch))

    # Contact table - V3 Services details matching competitor layout
    contact_data = [
        ['V3 Services Ltd', 'Tel:', '0203 576 1343'],
        ['117 Dartford Road', 'Email:', 'Info@V3-Services.com'],
        ['Dartford, DA1 3EN', '', ''],
    ]

    contact_table = Table(contact_data, colWidths=[2.5*inch, 0.7*inch, 2.5*inch])
    contact_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(contact_table)

    # ============ PAGE 2: CHARITY SUPPORT ============
    elements.append(PageBreak())

    # Add spacer to push content down from header
    elements.append(Spacer(1, 1.2*inch))

    # Page 2 Title
    elements.append(Paragraph("SUPPORT AND RESOURCES", styles['RSCharityTitle']))
    elements.append(Spacer(1, 0.1*inch))

    intro_text = """If you are experiencing homelessness or housing difficulties, the following organisations
    may be able to provide support and assistance:"""
    elements.append(Paragraph(intro_text, styles['RSCharityText']))
    elements.append(Spacer(1, 0.2*inch))

    # Charity logos directory
    assets_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'static', 'notice-assets')

    # Define charities with their details
    charities = [
        {
            'name': 'Shelter',
            'logo': 'shelter-logo.png',
            'phone': '0808 800 4444',
            'website': 'www.shelter.org.uk',
            'description': 'Housing advice and support'
        },
        {
            'name': 'Crisis',
            'logo': 'crisis-logo.png',
            'phone': '0300 636 1967',
            'website': 'www.crisis.org.uk',
            'description': 'Ending homelessness'
        },
        {
            'name': 'Mind',
            'logo': 'mind-logo.png',
            'phone': '0300 123 3393',
            'website': 'www.mind.org.uk',
            'description': 'Mental health support'
        },
        {
            'name': 'Samaritans',
            'logo': 'samaritans-logo.png',
            'phone': '116 123',
            'website': 'www.samaritans.org',
            'description': '24/7 emotional support'
        },
        {
            'name': 'The Big Issue Foundation',
            'logo': 'big-issue-logo.png',
            'phone': '020 7526 3200',
            'website': 'www.bigissue.org.uk',
            'description': 'Support for Big Issue vendors'
        },
    ]

    # Create charity entries
    for charity in charities:
        charity_section = []

        # Try to load logo
        logo_path = os.path.join(assets_dir, charity['logo'])
        if os.path.exists(logo_path):
            try:
                logo_img = Image(logo_path, width=1.2*inch, height=0.6*inch)
                logo_img.hAlign = 'CENTER'
                charity_section.append(logo_img)
            except Exception:
                charity_section.append(Paragraph(f"<b>{charity['name']}</b>", styles['RSSubtitle']))
        else:
            charity_section.append(Paragraph(f"<b>{charity['name']}</b>", styles['RSSubtitle']))

        # Charity details
        charity_section.append(Paragraph(charity['description'], styles['RSCharityText']))
        charity_section.append(Paragraph(f"<b>Phone:</b> {charity['phone']}", styles['RSCharityText']))
        charity_section.append(Paragraph(f"<b>Website:</b> {charity['website']}", styles['RSCharityText']))
        charity_section.append(Spacer(1, 0.15*inch))

        elements.extend(charity_section)

    elements.append(Spacer(1, 0.2*inch))

    # Local Authority Section
    local_authority = data.get('local_authority', '')
    if local_authority:
        elements.append(Paragraph("<b>Local Authority Housing Services</b>", styles['RSSubtitle']))
        elements.append(Paragraph(f"Contact {local_authority} Housing Department for assistance with housing.", styles['RSCharityText']))
        elements.append(Spacer(1, 0.1*inch))

    # Final message
    final_text = """These organisations provide free, confidential support. Please reach out if you need help."""
    elements.append(Paragraph(final_text, styles['RSCharityText']))

    return elements


@notices_bp.route('/admin/notices/types', methods=['GET'])
@jwt_required()
def get_notice_types():
    """Get available notice types"""
    user = require_admin()
    if not user:
        return jsonify({'error': 'Forbidden'}), 403
    
    notice_types = [
        {
            'id': 'notice_to_vacate',
            'name': 'Notice to Vacate',
            'description': 'Legal notice for squatters to vacate commercial property',
            'fields': [
                {'name': 'property_address', 'label': 'Property Address', 'type': 'textarea', 'required': True},
                {'name': 'client_name', 'label': 'Client Name', 'type': 'text', 'required': True},
                {'name': 'date', 'label': 'Date Notice Served', 'type': 'date', 'required': True},
                {'name': 'contact_phone', 'label': 'Contact Phone', 'type': 'text', 'required': False, 'default': '0203 576 1343'},
                {'name': 'contact_email', 'label': 'Contact Email', 'type': 'email', 'required': False, 'default': 'Info@V3-Services.com'},
            ]
        },
        {
            'id': 'abandoned_vehicle',
            'name': 'Abandoned Vehicle Report',
            'description': 'Report for abandoned vehicles on client property',
            'fields': [
                {'name': 'location', 'label': 'Location', 'type': 'textarea', 'required': True},
                {'name': 'client_name', 'label': 'Client Name', 'type': 'text', 'required': True},
                {'name': 'registration', 'label': 'Vehicle Registration', 'type': 'text', 'required': True},
                {'name': 'make_model', 'label': 'Make/Model', 'type': 'text', 'required': True},
                {'name': 'colour', 'label': 'Colour', 'type': 'text', 'required': True},
                {'name': 'condition', 'label': 'Vehicle Condition', 'type': 'text', 'required': True},
                {'name': 'description', 'label': 'Description', 'type': 'textarea', 'required': True},
                {'name': 'agent_name', 'label': 'Agent Name', 'type': 'text', 'required': True},
                {'name': 'date', 'label': 'Report Date', 'type': 'date', 'required': True},
            ]
        },
        {
            'id': 'rough_sleeper',
            'name': 'Rough Sleeper Notice',
            'description': 'Notice for rough sleepers with charity support information. Downloads as Word document for editing before printing as PDF.',
            'output_format': 'docx',
            'fields': [
                {'name': 'property_address', 'label': 'Property Address', 'type': 'textarea', 'required': True},
                {'name': 'landowner_name', 'label': 'Landowner Name', 'type': 'text', 'required': True},
                {'name': 'date_served', 'label': 'Date Served', 'type': 'date', 'required': True},
                {'name': 'vacate_date', 'label': 'Vacate By Date', 'type': 'date', 'required': True},
                {'name': 'vacate_time', 'label': 'Vacate By Time', 'type': 'text', 'required': False, 'default': '12:00'},
                {'name': 'local_authority', 'label': 'Local Authority (optional)', 'type': 'text', 'required': False},
            ]
        }
    ]
    
    return jsonify(notice_types)

@notices_bp.route('/admin/notices/generate', methods=['POST'])
@jwt_required()
def generate_notice():
    """Generate a notice document (PDF or Word depending on type)"""
    try:
        user = require_admin()
        if not user:
            current_app.logger.error("Access denied - user is not admin")
            return jsonify({'error': 'Forbidden'}), 403

        data = request.json
        if not data:
            current_app.logger.error("No JSON data received")
            return jsonify({'error': 'No data provided'}), 400

        notice_type = data.get('notice_type')
        notice_data = data.get('data', {})

        current_app.logger.info(f"Admin {user.email} generating notice type: {notice_type}")
        current_app.logger.info(f"Notice data: {notice_data}")

        # Rough sleeper uses Word template - returns .docx for editing
        if notice_type == 'rough_sleeper':
            docx_bytes = generate_rough_sleeper_docx(notice_data)
            docx_buffer = io.BytesIO(docx_bytes)
            docx_buffer.seek(0)
            filename = f"Rough_Sleeper_Notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            current_app.logger.info(f"Successfully generated {filename}")
            return send_file(
                docx_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )

        # Other notice types use PDF generation
        pdf_bytes = generate_notice_pdf(notice_type, notice_data)

        # Create buffer
        pdf_buffer = io.BytesIO(pdf_bytes)
        pdf_buffer.seek(0)

        # Create filename
        if notice_type == 'notice_to_vacate':
            filename = f"Notice_to_Vacate_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        elif notice_type == 'abandoned_vehicle':
            filename = f"Abandoned_Vehicle_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        else:
            filename = f"Notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

        current_app.logger.info(f"Successfully generated {filename}")

        return send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        current_app.logger.error(f"Error generating notice: {str(e)}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        return jsonify({'error': f'Failed to generate notice: {str(e)}'}), 500
